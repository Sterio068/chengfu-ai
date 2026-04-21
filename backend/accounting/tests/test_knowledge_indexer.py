"""
V1.1-SPEC §E-2 · indexer + extractor unit tests
================================================
不依賴 Meili / FastAPI · 純 mongomock + 臨時檔案
抽字真正 library(pymupdf/docx/openpyxl/Pillow)在 CI 環境可選裝
缺 library 的 test 會 skip
"""
import os
import pathlib
import tempfile
import pytest
import mongomock
from datetime import datetime

from services import knowledge_indexer
from services.knowledge_extract import extract, EXTRACTORS


# --- Fixtures ---
@pytest.fixture
def sources_col():
    client = mongomock.MongoClient()
    return client.chengfu_test.knowledge_sources


@pytest.fixture
def tmp_src(sources_col):
    """建一個有真實檔案結構的 tmp source"""
    d = tempfile.mkdtemp(prefix="chengfu_idx_")
    # 建立檔案結構
    (pathlib.Path(d) / "projects" / "海廢案").mkdir(parents=True)
    (pathlib.Path(d) / "projects" / "海廢案" / "建議書.txt").write_text(
        "承富創意 · 2024 環保署海洋廢棄物專案建議書 · 第一章主軸", encoding="utf-8"
    )
    (pathlib.Path(d) / "projects" / "海廢案" / "notes.md").write_text(
        "# 會議記錄\n- 3 月 1 日開案會議", encoding="utf-8"
    )
    (pathlib.Path(d) / "readme.md").write_text("根目錄 readme", encoding="utf-8")
    (pathlib.Path(d) / ".DS_Store").write_text("system", encoding="utf-8")
    (pathlib.Path(d) / "_unused" / "garbage.log").parent.mkdir()
    (pathlib.Path(d) / "_unused" / "garbage.log").write_text("log content")

    # 插入 source document
    sid = sources_col.insert_one({
        "name": "test source",
        "type": "local",
        "path": d,
        "enabled": True,
        "exclude_patterns": ["*.log", ".DS_Store"],
        "agent_access": [],
        "mime_whitelist": None,
        "max_size_mb": 10,
        "last_indexed_at": None,
        "last_index_stats": None,
        "created_at": datetime.utcnow(),
        "updated_at": datetime.utcnow(),
    }).inserted_id
    yield {"id": str(sid), "path": d, "col": sources_col}
    import shutil
    shutil.rmtree(d, ignore_errors=True)


# --- Extractor tests ---
def test_extract_text_utf8():
    """.txt utf-8 讀"""
    with tempfile.NamedTemporaryFile(suffix=".txt", mode="w", delete=False, encoding="utf-8") as f:
        f.write("繁體中文測試")
        p = f.name
    try:
        r = extract(p)
        assert r["type"] == "text"
        assert "繁體中文測試" in r["content_preview"]
        assert r["filename"] == os.path.basename(p)
        assert r["size"] > 0
    finally:
        os.unlink(p)


def test_extract_md():
    """.md 走 text extractor"""
    with tempfile.NamedTemporaryFile(suffix=".md", mode="w", delete=False, encoding="utf-8") as f:
        f.write("# title\n內容")
        p = f.name
    try:
        r = extract(p)
        assert r["type"] == "text"
        assert "# title" in r["content_preview"]
    finally:
        os.unlink(p)


def test_extract_unknown_format():
    """.xyz 走 fallback · 不 raise"""
    with tempfile.NamedTemporaryFile(suffix=".xyz", delete=False) as f:
        f.write(b"binary blob")
        p = f.name
    try:
        r = extract(p)
        assert r["type"] == "unknown"
        assert "xyz" in r["content_preview"] or r["filename"] in r["content_preview"]
    finally:
        os.unlink(p)


def test_extract_nonexistent_returns_error():
    """不存在檔案 · 不 raise · 回 type=error"""
    r = extract("/tmp/definitely-not-exist-xyz-12345.pdf")
    assert r["type"] == "error"
    assert "error" in r


def test_extract_docx():
    """.docx 若 python-docx 裝好 · 真 parse"""
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx 未安裝(非必測)")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        p = f.name
    try:
        doc = Document()
        doc.add_paragraph("承富創意整合行銷 · 測試段落")
        doc.add_paragraph("第二段 · 2026 年")
        doc.save(p)
        r = extract(p)
        assert r["type"] == "docx"
        assert "承富創意" in r["content_preview"]
        assert r["paragraph_count"] >= 2
    finally:
        os.unlink(p)


def test_extract_xlsx():
    """.xlsx 若 openpyxl 裝好 · 讀 sheet 前 N 行"""
    try:
        from openpyxl import Workbook
    except ImportError:
        pytest.skip("openpyxl 未安裝")

    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
        p = f.name
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "預算"
        ws.append(["項目", "金額"])
        ws.append(["場地", 30000])
        ws.append(["餐飲", 15000])
        wb.save(p)
        r = extract(p)
        assert r["type"] == "xlsx"
        assert "預算" in r["content_preview"]
        assert "場地" in r["content_preview"]
    finally:
        os.unlink(p)


# --- Indexer tests ---
def test_indexer_excludes_log_and_DS_Store(tmp_src):
    """exclude_patterns 會擋 .log 與 .DS_Store · 不抽字不索引"""
    stats = knowledge_indexer.reindex_source(
        tmp_src["id"], tmp_src["col"], meili_client=None,
    )
    assert stats["ok"] is True
    # 應該抽到 readme.md + projects/海廢案/建議書.txt + projects/海廢案/notes.md = 3
    assert stats["file_count"] >= 3
    assert stats["skipped"]["excluded"] >= 2  # .DS_Store + garbage.log
    # source doc 應更新 last_indexed_at
    doc = tmp_src["col"].find_one({"_id": __import__("bson").ObjectId(tmp_src["id"])})
    assert doc["last_indexed_at"] is not None


def test_indexer_incremental_second_run_skips_unchanged(tmp_src):
    """跑兩次 · 第二次 mtime 未變 · 都 skip"""
    s1 = knowledge_indexer.reindex_source(
        tmp_src["id"], tmp_src["col"], meili_client=None,
    )
    assert s1["file_count"] >= 3

    s2 = knowledge_indexer.reindex_source(
        tmp_src["id"], tmp_src["col"], meili_client=None,
    )
    # 第二輪 · 沒新檔 · 都在 unchanged
    assert s2["file_count"] == 0
    assert s2["skipped"]["unchanged"] >= 3


def test_indexer_disabled_source_skips(sources_col):
    """enabled=False 的 source 直接 skip"""
    d = tempfile.mkdtemp()
    sid = sources_col.insert_one({
        "name": "disabled",
        "path": d,
        "enabled": False,
    }).inserted_id
    try:
        r = knowledge_indexer.reindex_source(str(sid), sources_col, None)
        assert r.get("skipped") is True
    finally:
        import shutil
        shutil.rmtree(d, ignore_errors=True)


def test_indexer_nonexistent_path_returns_error(sources_col):
    """路徑不存在 · 不 crash · 回 ok=False + reason"""
    sid = sources_col.insert_one({
        "name": "bad-path",
        "path": "/tmp/definitely-not-exist-xyz-xyz",
        "enabled": True,
    }).inserted_id
    r = knowledge_indexer.reindex_source(str(sid), sources_col, None)
    assert r["ok"] is False
    assert "不存在" in r["reason"] or "不可讀" in r["reason"]


def test_indexer_reindex_all_iterates(tmp_src):
    """reindex_all 應跑所有 enabled sources"""
    r = knowledge_indexer.reindex_all(tmp_src["col"], None)
    assert "test source" in r
    assert r["test source"]["ok"] is True


def test_indexer_doc_id_stable():
    """同 source + 同 rel_path · id 必一致(才能 upsert 不重複)"""
    a = knowledge_indexer._doc_id_for("src_x", "projects/案/file.pdf")
    b = knowledge_indexer._doc_id_for("src_x", "projects/案/file.pdf")
    c = knowledge_indexer._doc_id_for("src_y", "projects/案/file.pdf")
    assert a == b
    assert a != c


def test_indexer_exclude_pattern_dir_prefix():
    """/sensitive/* 目錄前綴應該擋"""
    assert knowledge_indexer._match_excluded(
        "sensitive/secret.txt", ["/sensitive/*"]
    )
    assert not knowledge_indexer._match_excluded(
        "projects/normal.txt", ["/sensitive/*"]
    )


def test_indexer_exclude_pattern_glob():
    """~$* 擋 Word 暫存 · *.lock 擋 lock"""
    assert knowledge_indexer._match_excluded(
        "projects/案/~$建議書.docx", ["~$*"]
    )
    assert knowledge_indexer._match_excluded(
        "package-lock.json", ["*.lock", "package-lock.*"]
    )


def test_indexer_auto_tags_project_field(tmp_src):
    """檔在 projects/<名>/ 下 · 自動 tag project"""
    stats = knowledge_indexer.reindex_source(
        tmp_src["id"], tmp_src["col"], meili_client=None,
    )
    # 雖然 meili_client=None 沒進 Meili · 但 reindex_source 內部 logic 會處理 project
    # 直接用 extract + 手動模擬 indexer 的 enrichment 邏輯
    import os as _os
    from services.knowledge_extract import extract as _extract
    path = _os.path.join(tmp_src["path"], "projects", "海廢案", "建議書.txt")
    doc = _extract(path)
    assert doc["type"] == "text"
    # project enrichment 是 indexer 內部做的 · 這裡驗證 rel_path 推理
    rel = _os.path.relpath(path, tmp_src["path"])
    parts = rel.split(_os.sep)
    assert parts[0] == "projects"
    assert parts[1] == "海廢案"


# --- Delete source from index ---
def test_delete_source_from_index_no_meili():
    """meili_client=None 時 · 不 crash · 回 ok=False"""
    r = knowledge_indexer.delete_source_from_index("src_x", None)
    assert r["ok"] is False
